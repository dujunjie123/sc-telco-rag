# -*- coding: utf-8 -*- 
"""
Project: Telco-RAG
Creator: AI
Create time: 2024-06-26 15:21
IDE: PyCharm
Introduction:
"""
import subprocess

from modelscope.outputs import OutputKeys
from modelscope.pipelines import pipeline
from modelscope.utils.constant import Tasks
import torch
from langchain_text_splitters import RecursiveCharacterTextSplitter, HTMLHeaderTextSplitter
from xinference.client import Client
import pickle
from doc_embedding import Storage
from Source.chunking import chunk_doc

from docx import Document
import os
from tqdm import tqdm


device = "cuda:7" if torch.cuda.is_available() else "cpu"
torch.set_default_device(device)


def ali_segmentation_english():
    p = pipeline(
        task=Tasks.document_segmentation,
        model='/data/LLMs/modelscop/iic/nlp_bert_document-segmentation_english-base')

    result = p(
        documents="""Many references [10-14] have shown that AI/ML inference for image processing with 
        device-network synergy can alleviate the pressure of computation, memory footprint, storage, 
        power and required data rate on devices, reduce end-to-end latency and energy consumption, and 
        improve the end-to-end accuracy, efficiency and privacy when compared to the local execution 
        approach on either side. The scheme of split AI/ML image recognition can be depicted in Figure 
        5.1.1-1. The CNN is split into two parts according to the current image recognition task and environment.
         The intention is to offload the computation-intensive, energy-intensive parts to network server, 
         whereas leave the privacy-sensitive and delay- sensitive parts at the end device. The device executes 
         the inference up to a specific CNN layer and sends the intermediate data to the network server. 
         The network server runs through the remaining CNN layers. While the model is developed or invocated, 
         the split AI/ML operation is based on the legacy model.""")
    print(result[OutputKeys.TEXT])


def recursive_spliter_test1():
    from langchain_text_splitters import RecursiveCharacterTextSplitter

    state_of_the_union = """Many references [10-14] have shown that AI/ML inference for image processing with device-network synergy can alleviate the pressure of computation, memory footprint, storage, power and required data rate on devices, reduce end-to-end latency and energy consumption, and improve the end-to-end accuracy, efficiency and privacy when compared to the local execution approach on either side. The scheme of split AI/ML image recognition can be depicted in Figure 5.1.1-1. The CNN is split into two parts according to the current image recognition task and environment. The intention is to offload the computation-intensive, energy-intensive parts to network server, whereas leave the privacy-sensitive and delay- sensitive parts at the end device. The device executes the inference up to a specific CNN layer and sends the intermediate data to the network server. The network server runs through the remaining CNN layers. While the model is developed or invocated, the split AI/ML operation is based on the legacy model."""

    text_splitter = RecursiveCharacterTextSplitter(
        # 设置一个非常小的块大小，仅作展示。
        chunk_size=200,
        chunk_overlap=20,
        length_function=len,
        is_separator_regex=False,
    )

    texts = text_splitter.create_documents([state_of_the_union])

    print(texts)


def recursive_spliter_test2():
    from langchain_text_splitters import RecursiveCharacterTextSplitter

    state_of_the_union = """Many references [10-14] have shown that AI/ML inference for image processing with device-network synergy can alleviate the pressure of computation, memory footprint, storage, power and required data rate on devices, reduce end-to-end latency and energy consumption, and improve the end-to-end accuracy, efficiency and privacy when compared to the local execution approach on either side. The scheme of split AI/ML image recognition can be depicted in Figure 5.1.1-1. The CNN is split into two parts according to the current image recognition task and environment. The intention is to offload the computation-intensive, energy-intensive parts to network server, whereas leave the privacy-sensitive and delay- sensitive parts at the end device. The device executes the inference up to a specific CNN layer and sends the intermediate data to the network server. The network server runs through the remaining CNN layers. While the model is developed or invocated, the split AI/ML operation is based on the legacy model."""

    text_splitter = RecursiveCharacterTextSplitter(
        separators=["\n", ".", ",", ";", "!", "?"],
        chunk_size=200,
        chunk_overlap=20,
        length_function=len,
        is_separator_regex=False,
    )

    texts = text_splitter.create_documents([state_of_the_union])

    print(texts)


# recursive_spliter_test2()
def html_splitter1():
    from langchain_text_splitters import HTMLHeaderTextSplitter

    html_string = """
    <!DOCTYPE html>
    <html>
    <body>
        <div>
            <h1>标题1</h1>
            <p>关于标题1的一些介绍文本。</p>
            <div>
                <h2>子标题1</h2>
                <p>关于子标题1的一些介绍文本。</p>
                <h3>子子标题1</h3>
                <p>关于子子标题1的一些文本。</p>
                <h3>子子标题2</h3>
                <p>关于子子标题2的一些文本。</p>
            </div>
            <div>
                <h2>子标题2</h2>
                <p>关于子标题2的一些文本。</p>
            </div>
            <br>
            <p>关于标题1的一些结束文本。</p>
        </div>
    </body>
    </html>
    """

    headers_to_split_on = [
        ("h1", "一级标题"),
        ("h2", "二级标题"),
        ("h3", "三级标题"),
    ]

    html_splitter = HTMLHeaderTextSplitter(headers_to_split_on=headers_to_split_on)
    html_header_splits = html_splitter.split_text(html_string)
    print(html_header_splits)


def html_splitter2():
    from langchain_text_splitters import RecursiveCharacterTextSplitter, HTMLHeaderTextSplitter

    headers_to_split_on = [
        ("h1", "Header 1"),
        ("h2", "Header 2"),
        ("h3", "Header 3"),
        ("h4", "Header 4"),
    ]


    html_splitter = HTMLHeaderTextSplitter(headers_to_split_on=headers_to_split_on)

    # html_header_splits = html_splitter.split_text_from_url(url)
    html_header_splits = html_splitter.split_text_from_file("新文件 9.html")


    chunk_size = 500
    chunk_overlap = 30
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size, chunk_overlap=chunk_overlap
    )

    splits = text_splitter.split_documents(html_header_splits)

    print(splits)


class SegmentationTools:

    @staticmethod
    def html_splitter(html_file):
        headers_to_split_on = [
            ("h1", "Header 1"),
            ("h2", "Header 2"),
            ("h3", "Header 3"),
            ("h4", "Header 4"),
        ]

        html_splitter = HTMLHeaderTextSplitter(headers_to_split_on=headers_to_split_on)

        # html_header_splits = html_splitter.split_text_from_url(url)
        html_header_splits = html_splitter.split_text_from_file(html_file)

        return html_header_splits

    @staticmethod
    def split_by_recurs_separators(content, separators=None, chunk_size=500, chunk_overlap=30):

        if separators is None:
            separators = [".", "!", "?", ";", "\n", ","]

        text_splitter = RecursiveCharacterTextSplitter(
            separators=separators, chunk_size=chunk_size, chunk_overlap=chunk_overlap
        )

        splits = text_splitter.split_text(content)

        return splits

    @staticmethod
    def split_by_ali_segmentation_english_recurs_separators(content, separators=None, chunk_size=500, chunk_overlap=30):
        """
        使用阿里的语义切分+递归切分
        :param content:
        :param separators:
        :param chunk_size:
        :param chunk_overlap:
        :return:
        """

        if separators is None:
            separators = ["\n", ".", ",", ";", "!", "?"]

        text_splitter = RecursiveCharacterTextSplitter(
            separators=separators, chunk_size=chunk_size, chunk_overlap=chunk_overlap
        )

        splits = text_splitter.split_text(content)

        return splits


def get_doc_json(file_path='/data/data/game/itu/data/3GPP-Release18/Documents/22101-i50.docx',chunk_size=500):

    cont = []
    pos = []
    split = []

    obj = Document(file_path)

    for index, p in enumerate(obj.paragraphs):
        style_name = p.style.name
        if style_name.startswith('Heading'):
            # print(style_name,p.text,sep=':')
            pos.append(int(style_name[-1]))
            split.append(index)
            cont.append(p.text)

    split.append(len(obj.paragraphs))

    head_content = {}
    for index, head in enumerate(cont):
        head_content[head] = obj.paragraphs[split[index]:split[index + 1]]

    infos = {}
    relations = []
    for i in range(len(pos)):
        dic = {'title': cont[i]}
        infos[i] = dic

    # construct keys
    # input=[1, 2, 2, 1, 2, 2, 2, 1, 2, 2, 3, 3]
    # output = ['1', '1-1', '1-2', '2', '2-1', '2-2', '2-3', '3', '3-1', '3-2', '3-2-1', '3-2-2']
    keys = ['0'] * len(pos)

    def dfs(p, f):
        x = 1
        if keys[p] == '0':
            keys[p] = f + '-' + str(x)
        x += 1
        th = keys[p]
        for i in range(p + 1, len(pos)):
            if pos[i] - 1 == pos[p]:
                dfs(i, th + '-')
            elif pos[i] == pos[p]:
                if keys[i] == '0':
                    keys[i] = f + '-' + str(x)
                    x += 1
                    th = keys[i]
            elif pos[i] + 1 == pos[p]:
                return

    dfs(0, '')

    keys = [e.strip('-').replace('--', '-') for e in keys]

    def generate_tree(arr):
        top = current = {'id': 0, 'children': []}
        stack = [current]
        for pos, i in enumerate(arr):
            node = {'id': i, 'key': keys[pos], 'title': cont[pos], 'children': []}
            while i <= current['id']:
                stack.pop()
                current = stack[-1]
            current['children'].append(node)
            stack.append(node)
            current = node
        return top['children']

    arr = generate_tree(pos)
    # # json1 = json.loads(tree_dict, strict=False)
    # print(pos)
    # print(cont)
    # print(arr)
    records = []
    def combine_heads(head, p=""):
        if p:
            record = p + "\u0019" + head["title"]
        else:
            record = head["title"]
        records.append(record)
        if head["children"]:
            for c in head["children"]:
                combine_heads(c, p=record)

    for head in arr:
        combine_heads(head)


    head_structure_content = {}

    for record in records:
        if ("Foreword" in record or "Scope" in record or "References" in record or "Definition" in record or
                "abbreviation" in record or "Abbreviation" in record):
            continue

        contents = head_content[record.split("\u0019")[-1]]

        split_contents = []

        tmp = ""
        for content in contents:
            if content.style.name == "Normal" and len(tmp) != 0 and len(content.text) > 0 and content.text[-1] != ":":
                split_contents.append(SegmentationTools.split_by_recurs_separators(content.text))
                split_contents.append(SegmentationTools.split_by_recurs_separators(tmp))
                tmp = ""
            elif content.style.name == "Normal" and len(tmp) == 0 and len(content.text) > 0 and content.text[-1] != ":":
                split_contents.append(SegmentationTools.split_by_recurs_separators(content=content.text,
                                                                                   chunk_size=chunk_size))
            else:
                if "Heading" not in content.style.name:
                    tmp += content.text

        if len(tmp) != 0:
            split_contents.append(SegmentationTools.split_by_recurs_separators(tmp,separators=[".", ",", ";", "!", "?"]))

        head_structure_content[record] = split_contents

    return head_structure_content


def get_docx_2_latex(path="/data/data/game/itu/data/3GPP-Release18/", file_name="33513-i10.docx"):
    """

    :return:
    """

    cmd = f"""docker run --rm  -v {path}:/data --user $(id -u):$(id -g) pandoc/latex /data/Documents/{file_name} -o /data/Documents_latex/{file_name}.tex"""
    result = subprocess.run(cmd, shell=True)

    return result.stdout


class DocEmbedding:
    """
    构建文档的embedding
    """
    def __init__(self,xinferenc_url="http://localhost:9998",model_uid = "bge-m3-embedding-0620"):
        self.client = Client(xinferenc_url)
        self.emb_model = self.client.get_model(model_uid)

    def get_doc_embedding_origin(self,
                                 save_folder,
                                 filter_file_name,
                                 chunk_size
                                 ):
        """
        telco 中提供的原始方法
        """

        storage_name = 'Documents.db'
        dataset_name = "Standard"

        storage = Storage(f'/data/data/game/itu/data/3GPP-Release18/{storage_name}')

        for filename in tqdm(filter_file_name):
            print("------------------", filename)
            try:
                data_dict = storage.get_dict_by_id(dataset_name, filename)

                document_ds = chunk_doc(data_dict,chunk_size=chunk_size)
                for doc in document_ds:
                    """
                    doc {"text":"","source":'22101-i50.docx',"emb":[]}
                    """
                    max_retry = 5
                    retey = 0
                    while retey < max_retry:
                        try:
                            emb = self.emb_model.create_embedding(doc["text"])
                            doc["emb"] = emb["data"][0]["embedding"]
                            break

                        except Exception as e:
                            retry += 1
                            print("retry:", str(retry), doc)

                save_path = f"{save_folder}/{filename}.pkl"
                with open(save_path, "wb") as f:
                    pickle.dump(document_ds, f)

                print("embedding:", save_path)
            except Exception as e:
                print("error", e, filename)

    def get_doc_embedding(self,document_ds,filename,
                          save_folder="/data/data/game/itu/data/3GPP-Release18/embedding_by_bgem3_chunk500/"):
        """
        document_ds :[{"text":"","source":'22101-i50.docx',"emb":[]},{"text":"","source":'22101-i50.docx',"emb":[]},...]
        :return:
        """
        try:
            for doc in document_ds:
                """
                doc {"text":"","source":'22101-i50.docx',"emb":[]}
                """
                max_retry = 5
                retey = 0
                while retey < max_retry:
                    try:
                        emb = self.emb_model.create_embedding(doc["text"])
                        doc["emb"] = emb["data"][0]["embedding"]
                        break

                    except Exception as e:
                        retry += 1
                        print("retry:", str(retry), doc)

            save_path = f"{save_folder}/{filename}.pkl"
            with open(save_path, "wb") as f:
                pickle.dump(document_ds, f)
            print("embedding:", save_path)
        except Exception as e:
            print("error", e, filename)


def task_4_doc_chunking_embedding(folder_path,
                                  save_path,
                                  embedding_model_name,
                                  xinferenc_url,
                                  chunk_size=500):

    """
    把3gpps 文件转换成 json，然后再拼接每个段落和标题，最后embedding
    把key中的\u0019 转换成\t

    folder_path = "/data/data/game/itu/data/3GPP-Release18/Documents/"
    save_path = "/data/data/game/itu/data/3GPP-Release18/embedding_by_bgem3_structure_chunk500/"
    embedding_model_name = "bge-m3-sft-0718-qh9S6yJj"
    xinferenc_url="http://localhost:9998"

    :return:
    """

    doc_embedding = DocEmbedding(xinferenc_url=xinferenc_url, model_uid=embedding_model_name)
    for filename in tqdm(os.listdir(folder_path)):
        print("-----------------", filename)
        file_full_path = os.path.join(folder_path, filename)
        try:
            head_structure_content = get_doc_json(file_path=file_full_path, chunk_size=chunk_size)

            document_ds = []
            # document_ds: [{"text": "", "source": '22101-i50.docx', "emb": []},
            #               {"text": "", "source": '22101-i50.docx', "emb": []}, ...]
            for k,contents in head_structure_content.items():
                k = k.replace("\u0019", "")
                k = k.replace("\t", "")
                for content in contents:
                    for s in content:
                        tmp = {"text": k + " : " + s, "source": filename, "emb": []}
                        document_ds.append(tmp)

            doc_embedding.get_doc_embedding(document_ds, filename,save_folder=save_path)
        except Exception as e:
            print("-------------------------error-----------------------")
            print(e)
            print(file_full_path)


if __name__ == '__main__':

    task_4_doc_chunking_embedding(folder_path="/data/data/game/itu/data/3GPP-Release18/Documents/",
                                  save_path="/data/data/game/itu/data/3GPP-Release18/test_embedding_by_bgem3_structure_chunk500/",
                                  xinferenc_url="http://localhost:9998",
                                  embedding_model_name="bge-m3-0708-lnx3hyK8",
                                  chunk_size=500)

